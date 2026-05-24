"""
Generator dokumen DOCX: Analisis Etika Komputasi Islam
terhadap kode HTML, CSS, dan JavaScript pada repo fork-spk-wisata.

Output: docs/Analisis_Etika_Komputasi_Islam.docx (3 halaman)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, size=10, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    return p


def add_code_block(doc, code, font_size=8.5):
    """Tampilkan kode dalam paragraf monospace dengan latar abu-abu."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(font_size)
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rpr.append(rfonts)
    # Border + shading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "BBBBBB")
        b.set(qn("w:space"), "4")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


# ============================================================
# Mulai dokumen
# ============================================================
doc = Document()

# Margin lebih ramping agar setiap pembahasan muat 1 lembar
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)


# ============================================================
# HALAMAN 1 — HTML
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("Analisis Etika Komputasi Islam — HTML")
trun.bold = True
trun.font.size = Pt(13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("Studi Kasus: berkas index.html pada aplikasi React-Wisata")
srun.italic = True
srun.font.size = Pt(10)

add_paragraph(
    doc,
    "Berkas index.html merupakan titik masuk (entry point) aplikasi web React-Wisata. "
    "Walaupun sederhana, struktur ini memuat sejumlah keputusan teknis yang dapat ditinjau "
    "dari sudut pandang etika komputasi Islam, yaitu prinsip amanah (kepercayaan), "
    "kejujuran (sidq), kemaslahatan, dan itqan (kerapian dalam bekerja).",
    size=10,
)

add_heading(doc, "Cuplikan Kode", level=2)
html_code = (
    '<!doctype html>\n'
    '<html lang="en">\n'
    '  <head>\n'
    '    <meta charset="UTF-8" />\n'
    '    <meta name="viewport" content="width=device-width, initial-scale=1.0" />\n'
    '    <meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate" />\n'
    '    <meta http-equiv="Pragma" content="no-cache" />\n'
    '    <meta http-equiv="Expires" content="0" />\n'
    '    <title>react-wisata</title>\n'
    '  </head>\n'
    '  <body>\n'
    '    <div id="root"></div>\n'
    '    <script type="module" src="/src/main.jsx"></script>\n'
    '  </body>\n'
    '</html>'
)
add_code_block(doc, html_code)

add_heading(doc, "Analisis Etika", level=2)
add_bullet(
    doc,
    "Kejujuran informasi (sidq): meta tag Cache-Control dan Pragma mencegah peramban "
    "menyajikan data lama. Pengguna selalu menerima informasi wisata yang mutakhir, "
    "selaras dengan QS. Al-Mutaffifin (83): 1-3 yang melarang pengurangan hak orang lain "
    "termasuk dalam bentuk informasi yang menyesatkan.",
)
add_bullet(
    doc,
    "Inklusivitas dan kemaslahatan: meta viewport memastikan tampilan menyesuaikan "
    "perangkat pengguna sehingga aplikasi adil bagi pengguna dengan layar kecil maupun besar. "
    "Hal ini mencerminkan asas rahmatan lil 'alamin dalam pelayanan publik digital.",
)
add_bullet(
    doc,
    "Itqan (profesionalisme): struktur HTML rapi, hanya memuat elemen yang diperlukan, "
    "tidak menyertakan skrip pihak ketiga yang berisiko mengumpulkan data tanpa izin. "
    "Sesuai hadis riwayat Imam Thabrani bahwa Allah mencintai pekerjaan yang dikerjakan "
    "dengan itqan (rapi dan profesional).",
)
add_bullet(
    doc,
    "Catatan perbaikan: atribut lang=\"en\" kurang akurat karena konten aplikasi "
    "berbahasa Indonesia. Untuk menjaga kejujuran metadata sebaiknya diubah menjadi "
    "lang=\"id\" agar pembaca layar dan mesin pencari mengenali bahasa konten dengan benar.",
)

add_heading(doc, "Kesimpulan", level=2)
add_paragraph(
    doc,
    "Secara umum berkas index.html telah memenuhi prinsip etika komputasi Islam pada aspek "
    "kejujuran, kerapian, dan kemaslahatan, dengan satu catatan minor pada akurasi atribut "
    "bahasa yang perlu diselaraskan dengan konten sebenarnya.",
    size=10,
)

add_heading(doc, "Referensi", level=2)
add_paragraph(
    doc,
    "[1] Departemen Agama RI. Al-Qur'an dan Terjemahannya, QS. Al-Mutaffifin (83): 1-3.",
    size=9,
)
add_paragraph(
    doc,
    "[2] Salleh, M. S. (2003). Pengenalan Pembangunan Berteraskan Islam. Kuala Lumpur: "
    "Utusan Publications.",
    size=9,
)
add_paragraph(
    doc,
    "[3] World Wide Web Consortium. (2024). HTML Living Standard — The lang attribute. "
    "https://html.spec.whatwg.org/",
    size=9,
)
add_paragraph(
    doc,
    "[4] Mubarrak, H. (2017). Etika Profesi dalam Perspektif Islam. Banda Aceh: Bandar Publishing.",
    size=9,
)

add_page_break(doc)


# ============================================================
# HALAMAN 2 — CSS
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("Analisis Etika Komputasi Islam — CSS")
trun.bold = True
trun.font.size = Pt(13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("Studi Kasus: berkas src/index.css pada aplikasi React-Wisata")
srun.italic = True
srun.font.size = Pt(10)

add_paragraph(
    doc,
    "Berkas index.css memuat aturan gaya global aplikasi. Kode ini dipilih karena "
    "menunjukkan tiga keputusan etis yang relevan: pemilihan font sistem, normalisasi "
    "tata letak, dan adaptasi terhadap perangkat pengguna. Ketiganya dapat ditinjau "
    "melalui prinsip 'adl (keadilan), amanah (perlindungan data), serta menghindari "
    "israf (pemborosan).",
    size=10,
)

add_heading(doc, "Cuplikan Kode", level=2)
css_code = (
    ":root {\n"
    "  font-family: system-ui, Arial, sans-serif;\n"
    "}\n"
    "html, body, #root {\n"
    "  height: 100%;\n"
    "  margin: 0;\n"
    "  overflow-x: hidden;\n"
    "}\n"
    "* { box-sizing: border-box; }\n"
    "img { max-width: 100%; height: auto; }\n"
    "@media (max-width: 768px) {\n"
    "  html { font-size: 14px; }\n"
    "}\n"
    "@media (max-width: 480px) {\n"
    "  html { font-size: 13px; }\n"
    "}"
)
add_code_block(doc, css_code)

add_heading(doc, "Analisis Etika", level=2)
add_bullet(
    doc,
    "Amanah terhadap data pengguna: penggunaan system-ui sebagai font utama berarti "
    "peramban memakai font yang sudah terpasang di perangkat tanpa memuat font dari server "
    "pihak ketiga seperti Google Fonts. Dengan demikian alamat IP pengguna tidak terkirim "
    "ke layanan eksternal, sejalan dengan prinsip menjaga aurat data (privasi) yang "
    "merupakan amanah dalam QS. An-Nisa (4): 58.",
)
add_bullet(
    doc,
    "Keadilan ('adl) lintas perangkat: aturan responsive (img max-width: 100% dan media "
    "query) memastikan pengguna dengan layar 480px atau lebih kecil tetap memperoleh "
    "tampilan yang nyaman. Tidak ada pengguna yang dipinggirkan karena keterbatasan "
    "perangkatnya, mencerminkan QS. Al-Maidah (5): 8 tentang berlaku adil.",
)
add_bullet(
    doc,
    "Menghindari israf (pemborosan): kode CSS sangat ringkas, tidak memuat efek visual "
    "berlebihan, animasi yang boros baterai, atau aset besar. Sikap hemat sumber daya "
    "ini selaras dengan QS. Al-Isra (17): 26-27 yang mencela pemborosan.",
)
add_bullet(
    doc,
    "Ihsan (kenyamanan): aturan overflow-x: hidden mencegah scroll horizontal yang "
    "mengganggu, sedangkan box-sizing: border-box menjadikan perhitungan ukuran lebih "
    "intuitif sehingga tampilan tidak mudah berantakan.",
)
add_bullet(
    doc,
    "Catatan perbaikan: belum ada dukungan eksplisit terhadap prefers-color-scheme "
    "(mode gelap) maupun prefers-reduced-motion. Penambahan keduanya akan meningkatkan "
    "aspek ihsan dan inklusivitas bagi pengguna dengan sensitivitas cahaya atau gerakan.",
)

add_heading(doc, "Kesimpulan", level=2)
add_paragraph(
    doc,
    "Berkas index.css selaras dengan etika komputasi Islam pada aspek perlindungan privasi, "
    "keadilan tampilan, dan efisiensi sumber daya. Penambahan dukungan preferensi pengguna "
    "akan menyempurnakan aspek ihsan tanpa menambah kompleksitas berarti.",
    size=10,
)

add_heading(doc, "Referensi", level=2)
add_paragraph(
    doc,
    "[1] Departemen Agama RI. Al-Qur'an dan Terjemahannya, QS. An-Nisa (4): 58; "
    "QS. Al-Maidah (5): 8; QS. Al-Isra (17): 26-27.",
    size=9,
)
add_paragraph(
    doc,
    "[2] Mozilla Developer Network. (2024). Using system fonts and Privacy considerations. "
    "https://developer.mozilla.org/",
    size=9,
)
add_paragraph(
    doc,
    "[3] Hassan, M. K. (2016). Islamic Ethics in the Digital Age. Kuala Lumpur: IIUM Press.",
    size=9,
)
add_paragraph(
    doc,
    "[4] W3C Web Accessibility Initiative. (2023). Making Content Usable for People with "
    "Cognitive and Learning Disabilities. https://www.w3.org/WAI/",
    size=9,
)

add_page_break(doc)


# ============================================================
# HALAMAN 3 — JAVASCRIPT
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("Analisis Etika Komputasi Islam — JavaScript")
trun.bold = True
trun.font.size = Pt(13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("Studi Kasus: berkas src/utils/formatTanggal.js pada aplikasi React-Wisata")
srun.italic = True
srun.font.size = Pt(10)

add_paragraph(
    doc,
    "Berkas formatTanggal.js menyediakan fungsi utilitas untuk memformat tanggal ke "
    "dalam bahasa Indonesia. Walaupun pendek, fungsi ini menampilkan beberapa keputusan "
    "teknis yang dapat ditinjau dari prinsip itqan (profesionalisme), sidq (kejujuran "
    "perilaku program), serta kepekaan budaya yang dibingkai oleh nilai-nilai Islam.",
    size=10,
)

add_heading(doc, "Cuplikan Kode", level=2)
js_code = (
    'const HARI = ["Minggu","Senin","Selasa","Rabu","Kamis","Jumat","Sabtu"];\n'
    'const BULAN = [\n'
    '  "Januari","Februari","Maret","April","Mei","Juni",\n'
    '  "Juli","Agustus","September","Oktober","November","Desember",\n'
    '];\n'
    'export function formatTanggalIndonesia(date = new Date()) {\n'
    '  const hari    = HARI[date.getDay()];\n'
    '  const tanggal = date.getDate();\n'
    '  const bulan   = BULAN[date.getMonth()];\n'
    '  const tahun   = date.getFullYear();\n'
    '  return `${hari}, ${tanggal} ${bulan} ${tahun}`;\n'
    '}'
)
add_code_block(doc, js_code)

add_heading(doc, "Analisis Etika", level=2)
add_bullet(
    doc,
    "Itqan dan kesederhanaan: fungsi ini berupa pure function yang deterministik, "
    "yaitu masukan yang sama selalu menghasilkan keluaran yang sama tanpa efek samping. "
    "Karakter ini memudahkan pengujian dan pemeliharaan, sesuai hadis bahwa Allah "
    "mencintai amal yang dikerjakan dengan itqan (HR. Thabrani).",
)
add_bullet(
    doc,
    "Sidq pada perilaku program: fungsi tidak menyembunyikan operasi tersembunyi "
    "seperti pengiriman data ke server eksternal, pencatatan log, atau pembacaan "
    "informasi pengguna. Transparansi semacam ini sejajar dengan kejujuran "
    "(QS. Al-Ahzab (33): 70-71) yang diperintahkan kepada setiap muslim.",
)
add_bullet(
    doc,
    "Amanah keamanan: tidak ada dependensi eksternal yang dipanggil. Hal ini menekan "
    "risiko supply-chain attack dan menjaga amanah data pengguna karena seluruh logika "
    "dapat diaudit secara langsung.",
)
add_bullet(
    doc,
    "Kepekaan budaya dan rahmah: pemilihan nama hari dan bulan dalam bahasa Indonesia "
    "menghormati pengguna lokal. Pendekatan ini menampilkan rahmah (kasih sayang) pada "
    "konteks budaya pemakai, bukan memaksa format asing.",
)
add_bullet(
    doc,
    "Catatan perbaikan: aplikasi wisata yang banyak menyentuh pengguna muslim dapat "
    "diperkaya dengan opsi kalender Hijriyah, misalnya melalui Intl.DateTimeFormat "
    "dengan kalender islamic-umalqura. Penambahan ini akan memperluas inklusivitas "
    "tanpa mengorbankan kompatibilitas. Selain itu, parameter date sebaiknya divalidasi "
    "(misal isNaN(date.getTime())) agar tidak mengembalikan string undefined ketika "
    "menerima tanggal yang tidak valid — sebuah bentuk kejujuran teknis kepada pemanggil.",
)

add_heading(doc, "Kesimpulan", level=2)
add_paragraph(
    doc,
    "Fungsi formatTanggalIndonesia telah memenuhi prinsip itqan, sidq, dan amanah dalam "
    "etika komputasi Islam. Penambahan validasi input dan opsi kalender Hijriyah akan "
    "menjadikannya lebih sempurna dan inklusif bagi pengguna muslim.",
    size=10,
)

add_heading(doc, "Referensi", level=2)
add_paragraph(
    doc,
    "[1] Departemen Agama RI. Al-Qur'an dan Terjemahannya, QS. Al-Ahzab (33): 70-71.",
    size=9,
)
add_paragraph(
    doc,
    "[2] At-Thabrani, Sulaiman bin Ahmad. Al-Mu'jam al-Awsath. Hadis tentang itqan dalam pekerjaan.",
    size=9,
)
add_paragraph(
    doc,
    "[3] ECMA International. (2024). ECMAScript 2024 Language Specification. https://tc39.es/ecma262/",
    size=9,
)
add_paragraph(
    doc,
    "[4] MDN Web Docs. (2024). Intl.DateTimeFormat — calendar option (islamic-umalqura). "
    "https://developer.mozilla.org/",
    size=9,
)
add_paragraph(
    doc,
    "[5] Hashim, R. (2017). Islamic Perspectives on Information Technology Ethics. "
    "International Journal of Islamic Thought, 12(1).",
    size=9,
)


# Simpan
out_path = "/projects/sandbox/fork-spk-wisata/docs/Analisis_Etika_Komputasi_Islam.docx"
doc.save(out_path)
print(f"Dokumen tersimpan di: {out_path}")
